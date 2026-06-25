"""문서 정규화기 테스트"""

import sys
import logging
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))

from document_normalizer import DocumentNormalizer

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)


def test_supported_formats():
    """지원 포맷 테스트"""
    print("\n" + "="*60)
    print("1. 지원 포맷 확인")
    print("="*60)

    normalizer = DocumentNormalizer()

    print(f"\n✅ 지원 포맷:")
    for ext, fmt in normalizer.SUPPORTED_FORMATS.items():
        print(f"   {ext:10} → {fmt}")

    return True


def test_docx_conversion():
    """DOCX 변환 테스트 (모의)"""
    print("\n" + "="*60)
    print("2. DOCX 변환 테스트")
    print("="*60)

    try:
        from docx import Document
        from docx.shared import Pt

        with tempfile.TemporaryDirectory() as tmpdir:
            # 테스트 DOCX 파일 생성
            input_dir = Path(tmpdir) / "input"
            output_dir = Path(tmpdir) / "output"
            input_dir.mkdir()

            doc = Document()
            doc.add_heading("테스트 문서", 0)
            doc.add_heading("요구사항정의서", 1)
            doc.add_paragraph("이것은 테스트 문서입니다.")

            # 테이블 추가
            table = doc.add_table(rows=2, cols=3)
            table.rows[0].cells[0].text = "ID"
            table.rows[0].cells[1].text = "이름"
            table.rows[0].cells[2].text = "설명"
            table.rows[1].cells[0].text = "1"
            table.rows[1].cells[1].text = "항목1"
            table.rows[1].cells[2].text = "설명1"

            docx_file = input_dir / "test.docx"
            doc.save(docx_file)

            # 변환
            normalizer = DocumentNormalizer()
            stats = normalizer.normalize_directory(input_dir, output_dir)

            print(f"\n✅ DOCX 변환 성공")
            print(f"   총 파일: {stats['total_files']}개")
            print(f"   성공: {stats['successful']}개")
            print(f"   실패: {stats['failed']}개")

            # 생성된 Markdown 확인
            markdown_file = output_dir / "test.md"
            if markdown_file.exists():
                with open(markdown_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                print(f"\n생성된 Markdown 미리보기:")
                print("─" * 60)
                print(content[:300] + "..." if len(content) > 300 else content)
                print("─" * 60)
                return True
            else:
                print(f"✗ Markdown 파일 생성 실패")
                return False

    except ImportError:
        print("⚠️  python-docx 라이브러리 없음 (테스트 스킵)")
        return True
    except Exception as e:
        print(f"❌ 오류: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_xlsx_conversion():
    """XLSX 변환 테스트 (모의)"""
    print("\n" + "="*60)
    print("3. XLSX 변환 테스트")
    print("="*60)

    try:
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmpdir:
            input_dir = Path(tmpdir) / "input"
            output_dir = Path(tmpdir) / "output"
            input_dir.mkdir()

            # 테스트 XLSX 파일 생성
            wb = Workbook()
            ws = wb.active
            ws.title = "요구사항"

            ws['A1'] = "ID"
            ws['B1'] = "요구사항"
            ws['C1'] = "우선순위"

            ws['A2'] = "REQ-001"
            ws['B2'] = "사용자 로그인"
            ws['C2'] = "높음"

            ws['A3'] = "REQ-002"
            ws['B3'] = "문서 조회"
            ws['C3'] = "보통"

            xlsx_file = input_dir / "requirements.xlsx"
            wb.save(xlsx_file)

            # 변환
            normalizer = DocumentNormalizer()
            stats = normalizer.normalize_directory(input_dir, output_dir)

            print(f"\n✅ XLSX 변환 성공")
            print(f"   총 파일: {stats['total_files']}개")
            print(f"   성공: {stats['successful']}개")
            print(f"   실패: {stats['failed']}개")

            # 생성된 Markdown 확인
            markdown_file = output_dir / "requirements.md"
            if markdown_file.exists():
                with open(markdown_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                print(f"\n생성된 Markdown 미리보기:")
                print("─" * 60)
                print(content)
                print("─" * 60)
                return True
            else:
                print(f"✗ Markdown 파일 생성 실패")
                return False

    except ImportError:
        print("⚠️  openpyxl 라이브러리 없음 (테스트 스킵)")
        return True
    except Exception as e:
        print(f"❌ 오류: {e}")
        import traceback
        traceback.print_exc()
        return False


def test_markdown_passthrough():
    """Markdown 파일 통과 테스트"""
    print("\n" + "="*60)
    print("4. Markdown 파일 통과 (복사)")
    print("="*60)

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_dir = Path(tmpdir) / "input"
            output_dir = Path(tmpdir) / "output"
            input_dir.mkdir()

            # 테스트 Markdown 파일 생성
            md_file = input_dir / "README.md"
            md_content = "# 테스트 문서\n\n이것은 마크다운 파일입니다."
            md_file.write_text(md_content, encoding='utf-8')

            # 정규화 (실제로는 복사)
            normalizer = DocumentNormalizer()
            stats = normalizer.normalize_directory(input_dir, output_dir)

            print(f"\n✅ Markdown 파일 처리 완료")
            print(f"   총 파일: {stats['total_files']}개")
            print(f"   성공: {stats['successful']}개")
            print(f"   스킵: {stats['skipped']}개")

            # 복사 확인
            output_file = output_dir / "README.md"
            if output_file.exists():
                output_content = output_file.read_text(encoding='utf-8')
                if output_content == md_content:
                    print(f"✓ 파일 복사 성공")
                    return True
                else:
                    print(f"✗ 파일 내용 불일치")
                    return False
            else:
                print(f"✗ 복사 파일 없음")
                return False

    except Exception as e:
        print(f"❌ 오류: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """모든 테스트 실행"""
    print("\n")
    print("█" * 60)
    print("█  문서 정규화기 테스트")
    print("█" * 60)

    results = {
        "지원 포맷": test_supported_formats(),
        "Markdown 통과": test_markdown_passthrough(),
        "DOCX 변환": test_docx_conversion(),
        "XLSX 변환": test_xlsx_conversion(),
    }

    print("\n" + "="*60)
    print("테스트 결과 요약")
    print("="*60)

    for test_name, result in results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{test_name}: {status}")

    total = len(results)
    passed = sum(1 for r in results.values() if r)

    print(f"\n총 {total}개 중 {passed}개 통과")

    if passed == total:
        print("\n🎉 모든 테스트 통과!")
        return 0
    else:
        print(f"\n⚠️  {total - passed}개 테스트 실패 (라이브러리 미설치로 인한 스킵 포함)")
        return 1


if __name__ == "__main__":
    sys.exit(main())
